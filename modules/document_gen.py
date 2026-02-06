from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

class DocumentGenerator:
    @staticmethod
    def create_word_doc(text):
        doc = Document()
        
        # Style Configuration
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Header
        heading = doc.add_heading('LAUDO NEUROPSICOLÓGICO', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('') # Spacer

        # Parsing the text to create simplified formatted paragraphs
        # Assuming the text comes in blocks, we just flush it to the doc
        for paragraph in text.split('\n'):
            p = paragraph.strip()
            if p:
                if p.isupper() and len(p) < 50: # Likely a subtitle
                    h = doc.add_heading(p, level=1)
                else:
                    doc.add_paragraph(p)
        
        # Footer
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = "Gerado via Antigravity Clinical AI - Dr. Rodrigo Abel Gomes"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
